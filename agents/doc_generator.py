import os, sys, json, re, textwrap, time, html
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))
from utils import load_json, save_json, ensure_dir, now_iso, effective_decision
from kimi_client import KimiClient
from gdrive_uploader import upload_cv_cl, GDRIVE_AVAILABLE as GDRIVE_UPLOADER_AVAILABLE
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from email_notifier import send_email

KIMI_TIMEOUT = 30
# Written by main(), read by agents/email_notifier.py so the daily digest
# can announce and attach whatever was generated this run.
DOCS_MANIFEST = os.path.join("digests", "generated_docs_latest.json")

# What gets DELIVERED to the candidate, as opposed to what gets generated.
# He asked for .docx only (2026-08-24): the PDF is the copy an employer
# receives, but the one he opens is the one he needs to correct, and the
# system will not always be right. Both formats are still generated and both
# still go to Drive, so the send-ready PDF is one click away when he wants
# it. Set DELIVER_FORMATS to ".pdf,.docx" to get both attached again.
DELIVER_FORMATS = tuple(
    f.strip().lower() for f in os.environ.get("DELIVER_FORMATS", ".docx").split(",")
    if f.strip())

def _generate_summary(client, profile, title, company, description):
    role = profile.get("role", "")
    prompt = (
        f"Candidate: {profile['name']} -- {role}.\n"
        + (f"Seeking: {profile['target_role']}\n" if profile.get("target_role") else "")
        + (f"In his own words: {profile['summary']}\n" if profile.get("summary") else "")
        # Bullets, not just job titles. Starved of material this produced
        # keyword soup ("Skilled in Python, SQL, REST APIs") rather than a
        # summary that knows anything -- the same failure that made the cover
        # letter invent a mechanism it had not been given.
        + "\nExperience:\n"
        + "\n".join(
            f"- {exp.get('title', '')} @ {exp.get('company', '')}"
            + "".join(f"\n    * {b}" for b in (exp.get("bullets") or [])[:3])
            for exp in profile.get("experience", [])[:3])
        + f"\n\nSkills: {', '.join(profile.get('skills', {}).get('technical_default', []))}\n\n"
        f"Job to tailor for: {title} at {company}.\n"
        f"Description: {_clip(description, 1500)}\n\n"
        "Write the Profile Summary that sits at the top of his CV: 3-4 sentences, "
        "at most 70 words, showing why THIS job fits what he already does.\n\n"
        f'HARD RULE: his job title is exactly "{role}", and it is printed directly '
        "above this summary on the same page. Never restate it differently, never "
        "drop a word from it, and never describe him as something else -- not as a "
        '"Software Developer" for a software role, not as a "Data Engineer" for a '
        "data role. A summary that renames him contradicts the line above it and "
        "reads as a template nobody checked. Write about what he DOES and where he "
        "is heading; if the role must be named, use that exact string.\n\n"
        "Same voice as the rest of his documents: concrete rather than adjectival, "
        'a named system or a number rather than "skilled in". No keyword lists, no '
        "filler (passionate, proven track record, results-driven).\n\n"
        f'NEVER name the employer being applied to ("{company}"). This paragraph '
        "sits on a CV, not a cover letter: a summary that says \"seeking the X "
        "internship\" makes the document single-use, and if the wrong file is "
        "attached it tells a reader he wants to work somewhere else. The letter "
        "already argues for this specific job. Point at the KIND of work instead.\n\n"
        "Quote his metrics exactly as written above, including any hedge. "
        '"reducing manual data entry by ~40%" must not become "automated 40% of '
        'manual data entry" -- that is a different claim, and he has to defend '
        "the number he prints in an interview.\n\n"
        "Prefer finished work to work in progress. A migration still under way is "
        "a task, not a capability, and it competes for words with what actually "
        "distinguishes him.\n\n"
        "Use only facts given above; invent nothing. Return ONLY the summary text."
    )
    r = client.chat(
        [{"role": "user", "content": prompt}],
        max_tokens=200,
        response_format={"type": "text"},
    )
    if r is None:
        # Fallback derived from the profile: the old literal described a
        # "Data Analyst ... JavaScript, Power BI, NetSuite", which stopped
        # being true when the positioning changed, and it goes into a PDF
        # sent to employers.
        return (profile.get("summary")
                or f"{profile.get('role', '')}. "
                   f"Skills: {', '.join(profile.get('skills', {}).get('technical_default', [])[:8])}.")
    return r.strip()

def _clip(text: str, limit: int) -> str:
    """Truncates while KEEPING the paragraph breaks. textwrap.shorten collapses
    all whitespace, which is fine for a job description and exactly wrong for
    anything the model is asked to imitate the rhythm of."""
    text = (text or "").strip()
    return text if len(text) <= limit else text[:limit].rstrip() + "..."


def _load_text_model(path: str) -> str:
    """Reads an optional plain-text style model (config/cv_model.txt,
    config/cover_letter_model.txt). These are restored in CI from the
    CV_MODEL_B64 / CL_MODEL_B64 secrets and were, until now, never read by
    any code path."""
    try:
        with open(path, encoding="utf-8") as f:
            return f.read().strip()
    except (OSError, UnicodeDecodeError):
        return ""


def _generate_cover_letter(client, profile, title, company, location, description, score):
    job_desc = textwrap.shorten(description, width=600, placeholder="...")

    # Every fact below comes from the profile. Nothing about the candidate is
    # hardcoded here: the languages line used to read "DE (A2, learning)" as
    # a literal, so the moment he reaches B1 the CV would say B1 while every
    # cover letter for the same job still said A2 -- and this is text sent to
    # employers. Same reason the relocation sentence is now conditional.
    languages = profile.get("languages", "")
    target = profile.get("target_role", "")
    summary = profile.get("summary", "")
    projects = profile.get("projects", [])
    # The BULLETS, not just the titles. The prompt asks for a concrete example
    # of how he works, ideally one where he audited himself, and used to hand
    # over nothing but a project NAME to build it from. The model duly invented
    # the details: "a guardrail firing too late in the sequence", "rebuilt the
    # threshold layer", "re-ran the validation set". None of that happened, and
    # all of it went into a letter he would have to defend in an interview.
    # Forbidding invention cannot work while the facts are withheld, and the
    # true story is right here in the profile: an ad buried its "fluent German
    # required" clause past the excerpt window and scored 96/100. That is
    # better than anything the model made up.
    project_line = "\n".join(
        f"- {pr.get('title', '')}"
        + "".join(f"\n    * {b}" for b in (pr.get("bullets") or []))
        for pr in projects[:2]
    )
    relocation = profile.get("relocation_date", "")

    # Optional voice reference: config/cover_letter_model.txt, restored in CI
    # from the CL_MODEL_B64 secret. It was being restored and then ignored by
    # every code path -- now it steers the tone.
    model_letter = _load_text_model("config/cover_letter_model.txt")

    prompt = (
        f"Candidate: {profile['name']} ({profile.get('role', '')})\n"
        f"Address: {profile.get('address', '')}\n"
        f"LinkedIn: {profile.get('linkedin', '')}\n"
        + (f"Seeking: {target}\n" if target else "")
        + (f"In his own words: {summary}\n" if summary else "")
        + "\nExperience:\n"
        + "\n".join(f"- {e.get('title', '')} @ {e.get('company', '')}"
                    for e in profile.get("experience", []))
        + f"\n\nSkills: {', '.join(profile.get('skills', {}).get('technical_default', []))}\n"
        + (f"Projects (use THESE facts for any project example, and never invent "
           f"a mechanism, bug or fix that is not written here):\n{project_line}\n"
           if project_line else "")
        + (f"Languages: {languages}\n" if languages else "")
        + (f"Relocated to Switzerland: {relocation}\n" if relocation else "")
        + f"\nJob to apply: {title} at {company} ({location or 'Switzerland'})\n"
        f"Description: {job_desc}\n"
        f"Evaluation score: {score}/100\n\n"
        "Write a cover letter in English, five or six short paragraphs.\n\n"
        "Open with the story, never with an application sentence. He did not start in "
        "tech: he studied Law and worked at the Court of Justice before moving into "
        "software. That belongs in the FIRST paragraph. It is the image that makes a "
        "recruiter stop reading job boards for a second, and it is what makes the move "
        "into engineering read as deliberate rather than random. Do not demote it to a "
        "caveat further down.\n\n"
        "Then carry the path forward as a journey into what he does now, and why that "
        "points at THIS role -- use the 'Seeking' line for direction, not his current "
        "job title. Somewhere include a concrete example of HOW he works, preferably one "
        "where he audited his own work and found his own mistake. Close by asking for "
        "one specific thing.\n\n"
        "Voice:\n"
        "- Vary the rhythm. Long sentences carry the story, short ones land the point.\n"
        "- Contractions are fine. It should sound like a person talking, not a form.\n"
        "- Prefer a named system or a number to any adjective.\n\n"
        "Never write, in any wording:\n"
        "- an opening of the shape \"I am applying for the X role at Y\" or \"I am writing "
        "to express my interest\", and never the phrase \"What draws me to it is\";\n"
        "- a chronological roll-call of jobs (\"I work today as... Before that... "
        "Earlier...\"). That cadence is the clearest tell of a machine reading a CV out "
        "loud. Weave the experience into the narrative instead;\n"
        "- filler that signals a machine: passionate, excited to, thrilled, leverage, "
        "align with, deep dive, fast-paced, I believe my skills, perfect fit, proven "
        "track record, hit the ground running.\n\n"
        "Mention the language situation only if the posting raises it, and state it "
        "exactly as the Languages line above says -- never upgrade it.\n\n"
        "Rules: no invented facts, employers, dates, metrics or technologies -- use ONLY "
        "what is above. This includes TECHNICAL DETAIL: describe his own projects only in "
        "the terms given, and never invent a specific mechanism, bug or fix to make the "
        "story land better. He has to defend every sentence in an interview. "
        "No addresses, dates or signatures in the body. No em dashes, and no en dashes "
        "either: write a plain hyphen, a comma, or restructure the sentence. "
        "Return ONLY the letter body, starting at the first paragraph."
    )
    if model_letter:
        prompt += (
            "\n\nReference letter previously written by the candidate. Match its voice, "
            "rhythm and paragraph shape. Do NOT copy its sentences or reuse its "
            "company-specific details:\n---\n"
            # NOT textwrap.shorten: that collapses every newline into a space,
            # so the reference used to arrive as one run-on block -- the model
            # was told to match a rhythm that had just been erased.
            + _clip(model_letter, 3200)
            + "\n---"
        )
    r = client.chat(
        [{"role": "user", "content": prompt}],
        max_tokens=800,
        response_format={"type": "text"},
    )
    if r is None:
        # Offline fallback, used when the LLM call fails. It used to be a
        # fully written letter with facts baked in ("relocated in April
        # 2025", "German (A2 level)", "JavaScript, Power BI") that drifted
        # out of date the moment the CV changed -- and this text goes to
        # employers. It now says only what the profile says, and says less.
        # It also opened with "I am writing to apply for the X role at Y",
        # which is the single sentence the letter is supposed to avoid. The
        # profile summary already starts with the story in his own voice, so
        # the fallback leads with that and adds as little of its own as it can.
        top_experience = (profile.get("experience") or [{}])[0]
        employer = str(top_experience.get("company", "")).split("--")[0].strip()
        r = (
            "Dear Hiring Manager,\n\n"
            + (f"{profile.get('summary', '')}\n\n" if profile.get("summary") else "")
            + (f"That is the work I do today as {top_experience.get('title', '')}"
               + (f" at {employer}" if employer else "") + ", and it is why "
               f"{title} at {company} caught my attention.\n\n"
               if top_experience.get("title") else "")
            + (f"Languages: {profile.get('languages', '')}.\n\n"
               if profile.get("languages") else "")
            + "I would welcome a conversation about what your team is building and "
            "where I could contribute. My project work is on GitHub"
            + (f" at {profile.get('github', '')}" if profile.get("github") else "")
            + f".\n\nKind regards,\n{profile.get('name', '')}"
        )
    return r.strip()

def _order_skills_for_job(skills, description):
    """Puts the skills the posting actually asks for first.

    Reordering only: nothing is added, nothing is dropped, nothing is
    reworded. A recruiter reads the first line of SKILLS and stops, so
    leading with the stack the job names is worth real money -- and unlike
    rewriting experience bullets, it cannot introduce a claim he would have
    to defend.

    Matching is plain case-insensitive containment rather than a regex: skill
    names here are multi-character and concrete ("Docker", "GitHub Actions"),
    and a clever pattern is not worth the risk of a wrong match in a document
    sent to an employer.
    """
    text = str(description or "").lower()
    if not text:
        return list(skills)
    wanted, rest = [], []
    for skill in skills:
        name = str(skill).strip()
        # Two characters minimum: a one-letter "skill" would match anything.
        (wanted if len(name) >= 2 and name.lower() in text else rest).append(skill)
    return wanted + rest


def _format_employer(company: str) -> str:
    """"Name -- what it is" becomes "Name (what it is)".

    The profile stores the descriptor after a "--" separator, which
    job_evaluator splits on to get the bare company name. Rendering it
    verbatim put TWO "--" on one CV line:

        AI Software Engineer Intern -- netzdenker -- Swiss-based digital
        agency, DACH market | 06.2026 - Present

    where the structural separator and the descriptor separator are the same
    token, so a reader cannot tell which is which. Parentheses make the
    descriptor visibly subordinate, which is what it is -- and the descriptor
    earns its place: netzdenker is not a name a recruiter recognises.
    """
    name, _, descriptor = str(company or "").partition("--")
    name = name.strip()
    descriptor = descriptor.strip()
    return f"{name} ({descriptor})" if descriptor else name


def _role_keywords(title):
    title_lower = title.lower()
    if any(k in title_lower for k in ["data", "analyst", "business intelligence", "bi"]):
        return "data_focused"
    if any(k in title_lower for k in ["ai", "ml", "machine learning", "nlp"]):
        return "ai_focused"
    return "default"

# Import FPDF2
try:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos
    FPDF_AVAILABLE = True
except Exception:
    FPDF_AVAILABLE = False
    print("WARNING: fpdf2 not available; PDF generation disabled. Install: pip install fpdf2>=2.8.0")

# python-docx: the EDITABLE copy of the same two documents. The PDF is what
# gets sent -- fixed layout, nobody edits it by accident -- but the model will
# occasionally write a sentence the candidate wants to change, and he should
# not have to retype a letter to fix one line. Same content, same order.
try:
    import docx
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False
    print("WARNING: python-docx not available; editable copies disabled.")

# Characters the model reaches for that Latin-1 cannot represent. Order
# matters: these are substituted BEFORE the encode, because "ignore" DELETES
# whatever it cannot map. That bug was live and visible in a letter that went
# to an employer: every em dash became a double space, leaving six holes
# mid-sentence ("LLM APIs  Claude, Kimi  and ship them"). The replace() calls
# that were meant to prevent it ran AFTER the encode had already dropped the
# character, so they were dead code.
_PDF_SUBSTITUTIONS = {
    "\u2019": "'", "\u2018": "'", "\u201c": '"', "\u201d": '"',
    "\u2013": "-", "\u2014": "-", "\u2212": "-",
    "\u2026": "...", "\u00a0": " ", "\u2022": "-", "\u2192": "->",
    "\u20ac": "EUR", "\u2122": "(TM)",
}


def _safe_text(text):
    """Makes text safe for the Latin-1-only PDF core fonts.

    Substitutes first, then drops whatever is still unrepresentable (emoji,
    other scripts), so nothing the model routinely writes vanishes silently
    from a document that goes to an employer.
    """
    if not text:
        return ""
    for bad, good in _PDF_SUBSTITUTIONS.items():
        text = text.replace(bad, good)
    return text.encode("latin-1", "ignore").decode("latin-1")

def _docx_heading(document, text):
    """Section heading, matching the PDF's bold 11pt block headings."""
    para = document.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    para.paragraph_format.space_after = Pt(2)
    return para


def _docx_body(document, text, size=10, bold=False, space_after=2):
    para = document.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(space_after)
    return para


def cv_docx(profile, job, summary, path):
    """The CV as an editable .docx. Mirrors cv_pdf section for section.

    No _safe_text here: .docx is UTF-8, so the em dashes and accents that the
    PDF's Latin-1 core fonts cannot represent survive intact. That is a
    feature of the editable copy, not an oversight.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed")
    d = docx.Document()
    for section in d.sections:
        section.top_margin = section.bottom_margin = Inches(0.5)
        section.left_margin = section.right_margin = Inches(0.6)

    _docx_body(d, profile["name"], size=16, bold=True, space_after=0)
    for line in (f"{profile.get('role', '')} | {profile.get('location', '')}",
                 f"{profile.get('phone', '')} | {profile.get('email', '')} | {profile.get('linkedin', '')}",
                 f"Permit: {profile.get('permit', '')} | Notice: {profile.get('notice_period', '')}"):
        _docx_body(d, line, space_after=0)

    photo = profile.get("photo_path", "")
    if photo and os.path.exists(photo):
        try:
            d.add_picture(photo, width=Inches(1.4))
        except Exception:
            pass

    _docx_heading(d, "PROFILE SUMMARY")
    _docx_body(d, summary)

    role_key = _role_keywords(job.get("title", ""))
    skill_key = (f"technical_{role_key}" if f"technical_{role_key}" in profile["skills"]
                 else "technical_default")
    tech_skills = profile["skills"].get(skill_key, profile["skills"]["technical_default"])
    # Lead with what the posting actually asks for. Reordering only.
    tech_skills = _order_skills_for_job(tech_skills, job.get("description", ""))
    _docx_heading(d, "SKILLS")
    _docx_body(d, "Technical: " + ", ".join(tech_skills)
               + " | Communication: " + ", ".join(profile["skills"].get("communication", []))
               + " | Certifications: " + ", ".join(profile["skills"].get("certifications", [])))

    _docx_heading(d, "EXPERIENCE")
    for exp in profile["experience"]:
        _docx_body(d, f"{exp['title']} - {_format_employer(exp['company'])} | {exp['period']}",
                   bold=True, space_after=0)
        for b in exp["bullets"]:
            bullet = d.add_paragraph(b, style="List Bullet")
            bullet.paragraph_format.space_after = Pt(0)
            for run in bullet.runs:
                run.font.size = Pt(10)

    projects = profile.get("projects") or []
    if projects:
        _docx_heading(d, "PROJECTS")
        for pr in projects[:2]:
            heading = pr.get("title", "")
            if pr.get("url"):
                heading = f"{heading} - {pr['url']}"
            _docx_body(d, heading, bold=True, space_after=0)
            for b in pr.get("bullets", []):
                bullet = d.add_paragraph(b, style="List Bullet")
                bullet.paragraph_format.space_after = Pt(0)
                for run in bullet.runs:
                    run.font.size = Pt(10)

    _docx_heading(d, "EDUCATION")
    for edu in profile["education"]:
        _docx_body(d, f"{edu['degree']} | {edu['institution']} | {edu['period']}", space_after=0)

    _docx_heading(d, "LANGUAGES")
    _docx_body(d, profile.get("languages", ""))

    # Omitted entirely when empty, exactly like the PDF: a dangling heading
    # with nothing under it looks like a mistake on a CV sent to an employer.
    if str(profile.get("hobbies", "")).strip():
        _docx_heading(d, "HOBBIES & INTERESTS")
        _docx_body(d, profile["hobbies"])

    d.save(path)


def cl_docx(profile, letter, title, company, location, path):
    """The cover letter as an editable .docx. Mirrors cl_pdf."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed")
    d = docx.Document()
    for section in d.sections:
        section.top_margin = section.bottom_margin = Inches(0.8)
        section.left_margin = section.right_margin = Inches(0.9)

    _docx_body(d, now_iso().split("T")[0], space_after=8)
    for line in (profile["name"], profile["address"], profile["phone"],
                 profile["email"], profile["linkedin"]):
        _docx_body(d, line, space_after=0)
    _docx_body(d, "", space_after=6)
    _docx_body(d, company, space_after=0)
    _docx_body(d, location or "Switzerland", space_after=10)

    _docx_body(d, f"Re: Application for {title}", size=11, bold=True, space_after=8)

    for para in letter.split("\n\n"):
        para = para.strip()
        if para:
            _docx_body(d, para, size=11, space_after=8)

    _docx_body(d, "Kind regards,", size=11, space_after=0)
    _docx_body(d, profile["name"], size=11)
    d.save(path)


def cv_pdf(profile, job, summary, path):
    if not FPDF_AVAILABLE:
        raise RuntimeError("fpdf2 is not installed")
    pdf = FPDF()
    # Was auto=False: anything past the bottom of page 1 was silently dropped,
    # and this profile has four roles with bullets. A CV that loses its last
    # section without telling anyone is worse than a two-page CV.
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()
    m = 10
    pdf.set_margins(m, m, m)
    w = 210 - 2*m

    # Header with photo. The image occupies x=165..200, so the header text is
    # given a reduced width instead of the full page: it used to run under the
    # photo, and the profile summary started before the image ended.
    photo_w = 0
    if os.path.exists(profile.get("photo_path", "")):
        try:
            pdf.image(profile["photo_path"], x=165, y=10, w=35)
            photo_w = 40
        except Exception:
            photo_w = 0
    header_w = w - photo_w

    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(header_w, 8, _safe_text(profile["name"]), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 10)
    for line in (f"{profile.get('role', '')} | {profile.get('location', '')}",
                 f"{profile.get('phone', '')} | {profile.get('email', '')} | {profile.get('linkedin', '')}",
                 f"Permit: {profile.get('permit', '')} | Notice: {profile.get('notice_period', '')}"):
        pdf.multi_cell(header_w, 5, _safe_text(line), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    # Clear the photo before the first full-width block.
    if photo_w:
        pdf.set_y(max(pdf.get_y(), 10 + 35 * 5 / 4 + 3))
    pdf.ln(3)

    # AI-tailored summary
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 6, "PROFILE SUMMARY", ln=True)
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(w, 5, _safe_text(summary), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    # Skills
    role_key = _role_keywords(job.get("title", ""))
    skill_key = f"technical_{role_key}" if f"technical_{role_key}" in profile["skills"] else "technical_default"
    tech_skills = profile["skills"].get(skill_key, profile["skills"]["technical_default"])
    # Lead with what the posting actually asks for. Reordering only.
    tech_skills = _order_skills_for_job(tech_skills, job.get("description", ""))
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 6, "SKILLS", ln=True)
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(w, 5, _safe_text(
        "Technical: " + ", ".join(tech_skills) +
        " | Communication: " + ", ".join(profile["skills"].get("communication", [])) +
        " | Certifications: " + ", ".join(profile["skills"].get("certifications", []))
    ), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    # Experience
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 6, "EXPERIENCE", ln=True)
    pdf.set_font("Helvetica", "", 10)
    for exp in profile["experience"]:
        pdf.set_font("Helvetica", "B", 10)
        # multi_cell, not cell: cell() clips at the right margin, and these
        # lines are long ("Business Process & NetSuite Intern -- Gestora de
        # Inteligencia de Credito S.A. -- credit-intelligence bureau | ...").
        pdf.multi_cell(w, 5, _safe_text(
            f"{exp['title']} - {_format_employer(exp['company'])} | {exp['period']}"),
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 10)
        for b in exp["bullets"]:
            pdf.multi_cell(w, 4, _safe_text("  -- " + b), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)
    pdf.ln(1)

    # Projects. The profile has carried this section since the beginning and
    # only the cover letter ever read it, so the CV silently omitted the
    # strongest evidence on it: an unattended production system the candidate
    # audited and corrected himself. For someone moving into engineering from
    # another field, that is what the job titles cannot say.
    #
    # Placed after EXPERIENCE, not before it: his current role is the relevant
    # technical one, and pushing it down the page to lead with a side project
    # reads as having no job.
    projects = profile.get("projects") or []
    if projects:
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 6, "PROJECTS", ln=True)
        for pr in projects[:2]:
            pdf.set_font("Helvetica", "B", 10)
            heading = pr.get("title", "")
            if pr.get("url"):
                heading = f"{heading} -- {pr['url']}"
            pdf.multi_cell(w, 5, _safe_text(heading), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", "", 10)
            for b in pr.get("bullets", []):
                pdf.multi_cell(w, 4, _safe_text("  -- " + b),
                               new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(1)
        pdf.ln(1)

    # Education
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 6, "EDUCATION", ln=True)
    pdf.set_font("Helvetica", "", 10)
    for edu in profile["education"]:
        pdf.multi_cell(w, 5, _safe_text(f"{edu['degree']} | {edu['institution']} | {edu['period']}"),
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    # Languages
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 6, "LANGUAGES", ln=True)
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 5, _safe_text(profile.get("languages", "")), ln=True)
    pdf.ln(2)

    # Hobbies -- omitted entirely when empty, rather than leaving a dangling
    # heading with nothing under it on a CV sent to an employer.
    if str(profile.get("hobbies", "")).strip():
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 6, "HOBBIES & INTERESTS", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(w, 5, _safe_text(profile["hobbies"]), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.output(path)

def cl_pdf(profile, letter, title, company, location, path):
    if not FPDF_AVAILABLE:
        raise RuntimeError("fpdf2 is not installed")
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    m = 15
    pdf.set_margins(m, m, m)
    w = 210 - 2*m

    # Date + addresses
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 5, now_iso().split("T")[0], ln=True)
    pdf.ln(4)
    pdf.multi_cell(w, 5, _safe_text(f"{profile['name']}\n{profile['address']}\n{profile['phone']}\n{profile['email']}\n{profile['linkedin']}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)
    pdf.multi_cell(w, 5, _safe_text(f"{company}\n{location or 'Switzerland'}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(5)

    # Subject
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 6, _safe_text(f"Re: Application for {title}"), ln=True)
    pdf.ln(3)

    # Body
    pdf.set_font("Helvetica", "", 11)
    for para in letter.split("\n\n"):
        para = para.strip()
        if para:
            pdf.multi_cell(w, 6, _safe_text(para), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(2)

    # Signature
    pdf.ln(4)
    pdf.cell(0, 6, _safe_text("Kind regards,"), ln=True)
    pdf.ln(1)
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 6, _safe_text(profile["name"]), ln=True)

    pdf.output(path)

def _email_docs_to_candidate(folder, title, company, score, paths):
    """Mails the generated CV/CL to the candidate.

    Why this exists: the repo is public, so the PDFs can be neither committed
    nor uploaded as an Actions artifact (they carry full name, phone, personal
    e-mail, LinkedIn, permit status and photo), and the Google Drive upload
    fails with "Service Accounts do not have storage quota" until
    GDRIVE_REFRESH_TOKEN_B64 is set. Without this the documents live only on
    the runner and die with it. Recipient is always GMAIL_RECIPIENT -- the
    candidate -- never a recruiter.
    """
    sender = os.environ.get("GMAIL_SENDER")
    password = os.environ.get("GMAIL_APP_PASSWORD")
    recipient = os.environ.get("GMAIL_RECIPIENT")
    if not (sender and password and recipient):
        print("  [mail] GMAIL_* not configured -- CV/CL not mailed (they live only on this machine)")
        return False

    body = (
        f"<p>Tailored application materials for <b>{html.escape(str(title))}</b> "
        f"at <b>{html.escape(str(company))}</b> (match {html.escape(str(score))}).</p>"
        "<p>Attached: CV and cover letter. Review them before sending anything: "
        "nothing has been submitted to the employer.</p>"
    )
    # The subject is a mail HEADER, and title/company are scraped from
    # third-party job alerts: a newline in either would inject headers.
    # Collapse whitespace and cap the length.
    subject = " ".join(f"[Job Hunt] CV + cover letter: {title} @ {company}".split())[:180]
    ok = send_email(recipient, subject, body, sender, password, attachments=paths)
    print(f"  [mail] {'Sent to candidate' if ok else 'Send FAILED'}: {len(paths)} file(s)")
    return ok


def generate_docs_for_job(client, profile, ev: dict, gen_dir: str = "generated_docs",
                          mail: bool = True, force: bool = False) -> str | None:
    """Generates (and, if configured, uploads to Drive) the CV/CL for a
    single evaluation record. Shared by main() (daily batch, reads
    job_evaluations_latest.json) and agents/add_job.py (single manual
    evaluation) -- the two used to be on separate tracks entirely, so a
    manually-added job scoring APPLY never got tailored materials no
    matter how high it scored. Returns the output folder, or None on
    failure/skip."""
    score = ev.get("score") or 0  # ERROR evaluations carry score None
    # Only APPLY jobs get tailored materials -- the EFFECTIVE decision:
    # a hard-blocked or low-confidence (insufficient_info) job keeps its
    # high score visible but must never trigger automatic CV/CL generation.
    # The APPLY gate exists to stop the pipeline generating documents on its
    # own for a job it is not confident about. It was never meant to stop the
    # CANDIDATE: "scoring is input, not gospel -- the user can override any
    # decision". force=True is that override, for when he reads a REVIEW job
    # and decides it is worth applying to. Nothing is sent either way; the
    # documents go to him.
    if not force and effective_decision(ev) != "APPLY":
        return None

    job = ev.get("job", ev)
    title = job.get("title", "Job")
    company = job.get("company", "Company")
    location = job.get("location", "")
    desc = job.get("description", "")

    safe_name = re.sub(r"[^\w\-]", "_", f"{company}_{title}")[:60]
    folder = os.path.join(gen_dir, safe_name)
    ensure_dir(folder)

    print(f"[doc_generator] Generating for {title} @ {company} (score {score})")

    try:
        summary = _generate_summary(client, profile, title, company, desc)
        letter = _generate_cover_letter(client, profile, title, company, location, desc, score)
    except Exception as e:
        print(f"  [doc_generator] API error for {company} -- skipping ({type(e).__name__}: {str(e)[:120]})")
        return None
    time.sleep(1.5)

    if FPDF_AVAILABLE:
        cv_pdf(profile, job, summary, os.path.join(folder, f"CV_{safe_name}.pdf"))
        cl_pdf(profile, letter, title, company, location, os.path.join(folder, f"CL_{safe_name}.pdf"))
        # Editable twins. The PDF is the copy that gets sent; the .docx exists
        # so a sentence the model got wrong can be fixed without retyping the
        # whole document. A failure here must never cost the PDFs.
        if DOCX_AVAILABLE:
            try:
                cv_docx(profile, job, summary, os.path.join(folder, f"CV_{safe_name}.docx"))
                cl_docx(profile, letter, title, company, location,
                        os.path.join(folder, f"CL_{safe_name}.docx"))
            except Exception as e:
                print(f"  [docx] Editable copies failed (continuing): "
                      f"{type(e).__name__}: {str(e)[:120]}")
        save_json(os.path.join(folder, "ai_summary.json"), {"summary": summary, "letter": letter, "score": score})
        print(f"  Saved to {folder}/")

        drive_link = ""

        # Upload to Google Drive. The returned folder link, when there is
        # one, is what turns the digest's "attached to this e-mail" into a
        # download link -- which is the shape the candidate asked for.
        if GDRIVE_UPLOADER_AVAILABLE:
            try:
                uploaded = upload_cv_cl(folder, company, title) or {}
                drive_link = uploaded.get("folder_link", "")
            except Exception as e:
                print(f"  [GDrive] Upload failed (continuing): {e}")

        # Record the link where main() can find it when it builds the digest
        # manifest. ai_summary.json already travels with the folder, and
        # rewriting it is cheaper than threading a second return value
        # through a function add_job.py also calls.
        save_json(os.path.join(folder, "ai_summary.json"),
                  {"summary": summary, "letter": letter, "score": score,
                   "drive_link": drive_link})

        # Mail the PDFs to the candidate. Independent of Drive on purpose:
        # Drive is the one that has been silently failing, and a generated
        # document that reaches nobody is the same as no document.
        #
        # mail=False in the daily batch: sending one e-mail per APPLY job
        # buried the notice in a second inbox thread. main() records the
        # documents in a manifest instead and the digest e-mail, which the
        # candidate already reads once a day, announces and carries them.
        # add_job.py keeps mail=True: it runs alone, with no digest to ride.
        if mail:
            try:
                _email_docs_to_candidate(
                    folder, title, company, score,
                    sorted(os.path.join(folder, f) for f in os.listdir(folder)
                           if f.lower().endswith(DELIVER_FORMATS)),
                )
            except Exception as e:
                print(f"  [mail] Failed (continuing): {type(e).__name__}: {str(e)[:120]}")
    else:
        save_json(os.path.join(folder, "ai_summary.json"), {"summary": summary, "letter": letter, "score": score})
        print(f"  Saved JSON only (fpdf2 missing): {folder}/")

    return folder


def main():
    evals = load_json("digests/job_evaluations_latest.json")
    profile = load_json("config/candidate_profile.json")
    client = KimiClient()
    gen_dir = "generated_docs"
    ensure_dir(gen_dir)

    if not evals:
        print("No evaluations found."); return
    if not profile:
        print("config/candidate_profile.json missing or invalid (check CANDIDATE_PROFILE_B64 secret) -- skipping doc generation")
        return

    manifest = []
    for ev in evals:
        folder = generate_docs_for_job(client, profile, ev, gen_dir, mail=False)
        if not folder:
            continue
        job = ev.get("job", ev)
        # A Drive link means the digest offers a download instead of
        # attaching the PDFs. Empty (Drive unconfigured or failed) keeps the
        # attachment path, so the documents reach him either way.
        link = (load_json(os.path.join(folder, "ai_summary.json")) or {}).get("drive_link", "")
        manifest.append({
            "title": job.get("title", "Job"),
            "company": job.get("company", "Company"),
            "score": ev.get("score"),
            "folder": folder,
            "link": link,
            "files": sorted(os.path.join(folder, f) for f in os.listdir(folder)
                            if f.lower().endswith((".pdf", ".docx"))),
        })

    # The digest e-mail reads this and announces (and carries) the documents,
    # so the candidate gets one message a day instead of one per APPLY job.
    save_json(DOCS_MANIFEST, {"generated_at": now_iso(), "documents": manifest})

    # Say so out loud. This step used to print absolutely nothing when no job
    # reached APPLY, which is indistinguishable from a crash in the CI log.
    print(f"[doc_generator] {len(manifest)} job(s) got tailored documents "
          f"out of {len(evals)} evaluated (only APPLY qualifies). "
          f"Manifest: {DOCS_MANIFEST}")

if __name__ == "__main__":
    main()
